from langchain.tools import tool
from pathlib import Path
from docx import Document
import os
import copy

def check_file(file_path:str) -> str:
    path = Path(file_path)
    if path.exists():
        if path.suffix.lower() == '.docx':
            doc = Document(file_path)
            paragraphs = []
            for paragraph in doc.paragraphs:
                paragraphs.append(paragraph.text)
            return "\n".join(paragraphs)
    return

@tool("file_reader")
def read_file(file_path:str) -> str:
    """Reads a file and returns its contents as plain text.
    
    Accepts .txt and .docx file types.
    
    Args:
        file_path: The full path to the file to read.
    Returns:
        The contents of the file as a plain string.
    """
    path = Path(file_path)
    if path.exists():
        if path.suffix.lower() == '.txt':
            with open(file_path) as f:
                return(f.read())
        elif path.suffix.lower() == '.docx':
            return check_file(file_path)
    else:
        return("No file exists")
    return 

@tool('exporter')
def file_exporter(tailored_resume:str, tailored_cover_letter:str, job_description:str, company_name: str, job_title:str, interview_prep:str, winning_resume:str) -> str:
    """Tool for exporting all the files into their respective file types, and folder them into a single folder.
    
    Args:
        tailored_resume: The finalized resume.
        tailored_cover_letter: The generated cover letter.
        job_description: The full job description.
        company_name: The company's name.
        job_title: The name of the position.
        interview_prep: The full study guide for interview.
    Return:
        The structured folder with all the contents for the job application.
    """
    
    output_dir = Path(os.getenv('OUTPUT_DIR', '~/job_applications')).expanduser().resolve()
    
    folder_path = output_dir / company_name / job_title
    folder_path.mkdir(parents=True, exist_ok=True)
    
    resume_doc = Document(f'inputs/resumes/{winning_resume}')
    for par in resume_doc.paragraphs:
        par.clear()
    for line in tailored_resume.split('\n'):
        if line.strip():
            resume_doc.add_paragraph(line.strip())
    resume_doc.save(folder_path / 'resume.docx')
    
    cover_letter_doc = Document(f'inputs/cover_letter_template.docx')
    for cov_par in cover_letter_doc.paragraphs:
        cov_par.clear()
    for cov_line in tailored_cover_letter.split('\n'):
        if cov_line.strip():
            cover_letter_doc.add_paragraph(cov_line.strip())
    cover_letter_doc.save(folder_path / 'cover_letter.docx')
    
    (folder_path / 'job_description.txt').write_text(job_description)
    (folder_path / 'interview_prep.md').write_text(interview_prep)
    
    return "Folder successfully created along with contents"